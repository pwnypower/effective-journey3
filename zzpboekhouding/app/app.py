import os
import sqlite3
import smtplib
import io
import urllib.request
import json as _json
from calendar import monthrange
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from datetime import datetime, date, timedelta
from flask import (Flask, render_template, request, redirect, url_for,
                   flash, g, send_file, make_response, jsonify)
from dotenv import load_dotenv

load_dotenv()

try:
    from apscheduler.schedulers.background import BackgroundScheduler
    HAS_SCHEDULER = True
except ImportError:
    HAS_SCHEDULER = False

try:
    from docx import Document as DocxDoc
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from docxtpl import DocxTemplate
    HAS_DOCXTPL = True
except ImportError:
    HAS_DOCXTPL = False

TEMPLATE_PAD = os.environ.get(
    "TEMPLATE_PAD",
    os.path.join(os.path.dirname(__file__), "factuur_template.docx")
)


class IngressMiddleware:
    def __init__(self, wsgi_app):
        self.wsgi_app = wsgi_app

    def __call__(self, environ, start_response):
        ingress_path = environ.get("HTTP_X_INGRESS_PATH", "")
        if ingress_path:
            environ["SCRIPT_NAME"] = ingress_path
            path_info = environ.get("PATH_INFO", "")
            if path_info.startswith(ingress_path):
                environ["PATH_INFO"] = path_info[len(ingress_path):]
        return self.wsgi_app(environ, start_response)


app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-change-me")
app.wsgi_app = IngressMiddleware(app.wsgi_app)

DB_PATH = os.environ.get("DB_PATH", os.path.join(os.path.dirname(__file__), "boekhouding.db"))
PORT = int(os.environ.get("PORT", "5000"))

# Env-var defaults (used for initial DB seed only)
_ENV_NAAM = os.environ.get("BEDRIJFSNAAM", "Mijn ZZP Bedrijf")
_ENV_KVK = os.environ.get("KVK_NUMMER", "")
_ENV_BTW = os.environ.get("BTW_NUMMER", "")
_ENV_IBAN = os.environ.get("IBAN", "")
_ENV_BASE_URL = os.environ.get("APP_BASE_URL", "").rstrip("/")
_ENV_SMTP_HOST = os.environ.get("SMTP_HOST", "")
_ENV_SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))
_ENV_SMTP_USER = os.environ.get("SMTP_USER", "")
_ENV_SMTP_PASS = os.environ.get("SMTP_PASSWORD", "")
_ENV_SMTP_FROM = os.environ.get("SMTP_FROM", _ENV_SMTP_USER)
_ENV_SMTP_TLS = os.environ.get("SMTP_USE_TLS", "true").lower() == "true"
_ENV_MOLLIE = os.environ.get("MOLLIE_API_KEY", "")


# ─── Database ─────────────────────────────────────────────────────────────────

def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA foreign_keys = ON")
    return g.db


@app.teardown_appcontext
def close_db(exception=None):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def init_db():
    db = sqlite3.connect(DB_PATH)
    db.executescript("""
        CREATE TABLE IF NOT EXISTS klanten (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            naam TEXT NOT NULL,
            contactpersoon TEXT,
            email TEXT,
            telefoon TEXT,
            adres TEXT,
            postcode TEXT,
            plaats TEXT,
            land TEXT DEFAULT 'Nederland',
            kvk TEXT,
            btw_nummer TEXT
        );
        CREATE TABLE IF NOT EXISTS facturen (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            factuurnummer TEXT NOT NULL UNIQUE,
            klant_id INTEGER NOT NULL,
            factuurdatum TEXT NOT NULL,
            vervaldatum TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'concept',
            betaallink TEXT,
            mollie_payment_id TEXT,
            notities TEXT,
            korting REAL NOT NULL DEFAULT 0,
            korting_type TEXT NOT NULL DEFAULT 'pct',
            projectnummer TEXT,
            referentie TEXT,
            betalingstermijn INTEGER NOT NULL DEFAULT 30,
            FOREIGN KEY (klant_id) REFERENCES klanten (id)
        );
        CREATE TABLE IF NOT EXISTS factuurregels (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            factuur_id INTEGER NOT NULL,
            omschrijving TEXT NOT NULL,
            aantal REAL NOT NULL DEFAULT 1,
            eenheid TEXT NOT NULL DEFAULT 'st',
            prijs_per_stuk REAL NOT NULL,
            btw_percentage REAL NOT NULL DEFAULT 21,
            FOREIGN KEY (factuur_id) REFERENCES facturen (id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS uitgaven (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            datum TEXT NOT NULL,
            omschrijving TEXT NOT NULL,
            categorie TEXT,
            bedrag_excl_btw REAL NOT NULL,
            btw_percentage REAL NOT NULL DEFAULT 21,
            zakelijk_percentage REAL NOT NULL DEFAULT 100,
            leverancier TEXT,
            notities TEXT
        );
        CREATE TABLE IF NOT EXISTS producten (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            naam TEXT NOT NULL,
            beschrijving TEXT,
            prijs REAL NOT NULL DEFAULT 0,
            eenheid TEXT NOT NULL DEFAULT 'st'
        );
        CREATE TABLE IF NOT EXISTS instellingen (
            id INTEGER PRIMARY KEY DEFAULT 1,
            bedrijfsnaam TEXT NOT NULL DEFAULT '',
            adres TEXT NOT NULL DEFAULT '',
            postcode TEXT NOT NULL DEFAULT '',
            stad TEXT NOT NULL DEFAULT '',
            email TEXT NOT NULL DEFAULT '',
            telefoon TEXT NOT NULL DEFAULT '',
            kvk TEXT NOT NULL DEFAULT '',
            btwnummer TEXT NOT NULL DEFAULT '',
            iban TEXT NOT NULL DEFAULT '',
            betalingstermijn INTEGER NOT NULL DEFAULT 30,
            factuurprefix TEXT NOT NULL DEFAULT 'FAC',
            factuurvolgend INTEGER NOT NULL DEFAULT 1,
            alg_voorwaarden TEXT NOT NULL DEFAULT '',
            juridisch_voetnoot TEXT NOT NULL DEFAULT '',
            mollie_key_test TEXT NOT NULL DEFAULT '',
            mollie_key_live TEXT NOT NULL DEFAULT '',
            mollie_mode TEXT NOT NULL DEFAULT 'test',
            smtp_host TEXT NOT NULL DEFAULT '',
            smtp_port INTEGER NOT NULL DEFAULT 587,
            smtp_user TEXT NOT NULL DEFAULT '',
            smtp_password TEXT NOT NULL DEFAULT '',
            smtp_from TEXT NOT NULL DEFAULT '',
            smtp_use_tls INTEGER NOT NULL DEFAULT 1,
            app_base_url TEXT NOT NULL DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS periodieke_facturen (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            naam TEXT NOT NULL,
            klant_id INTEGER NOT NULL,
            interval TEXT NOT NULL DEFAULT 'maandelijks',
            volgende_datum TEXT NOT NULL,
            actief INTEGER NOT NULL DEFAULT 1,
            betalingstermijn INTEGER NOT NULL DEFAULT 30,
            notities TEXT,
            FOREIGN KEY (klant_id) REFERENCES klanten (id)
        );
        CREATE TABLE IF NOT EXISTS periodieke_regels (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            periodiek_id INTEGER NOT NULL,
            omschrijving TEXT NOT NULL,
            aantal REAL NOT NULL DEFAULT 1,
            eenheid TEXT NOT NULL DEFAULT 'st',
            prijs_per_stuk REAL NOT NULL,
            btw_percentage REAL NOT NULL DEFAULT 21,
            FOREIGN KEY (periodiek_id) REFERENCES periodieke_facturen (id) ON DELETE CASCADE
        );
    """)

    # Column migrations (safe to run multiple times)
    for col, defn in [("telefoon", "TEXT"), ("betalingstermijn", "INTEGER")]:
        try:
            db.execute(f"ALTER TABLE klanten ADD COLUMN {col} {defn}")
        except Exception:
            pass
    for col, defn in [
        ("korting", "REAL NOT NULL DEFAULT 0"),
        ("korting_type", "TEXT NOT NULL DEFAULT 'pct'"),
        ("projectnummer", "TEXT"),
        ("referentie", "TEXT"),
        ("betalingstermijn", "INTEGER NOT NULL DEFAULT 30"),
    ]:
        try:
            db.execute(f"ALTER TABLE facturen ADD COLUMN {col} {defn}")
        except Exception:
            pass
    for col, defn in [("eenheid", "TEXT NOT NULL DEFAULT 'st'"),
                      ("periode_van", "TEXT"), ("periode_tot", "TEXT")]:
        try:
            db.execute(f"ALTER TABLE factuurregels ADD COLUMN {col} {defn}")
        except Exception:
            pass
    for col, defn in [("periode_van", "TEXT"), ("periode_tot", "TEXT")]:
        try:
            db.execute(f"ALTER TABLE periodieke_regels ADD COLUMN {col} {defn}")
        except Exception:
            pass
    for col, defn in [("mollie_poll_url", "TEXT NOT NULL DEFAULT ''"),
                      ("mollie_poll_token", "TEXT NOT NULL DEFAULT ''"),
                      ("mollie_relay_url", "TEXT NOT NULL DEFAULT ''"),
                      ("mollie_bevestiging_url", "TEXT NOT NULL DEFAULT ''")]:
        try:
            db.execute(f"ALTER TABLE instellingen ADD COLUMN {col} {defn}")
        except Exception:
            pass

    db.commit()

    # Seed settings from env vars on first run
    if not db.execute("SELECT id FROM instellingen WHERE id=1").fetchone():
        db.execute(
            "INSERT INTO instellingen (id,bedrijfsnaam,kvk,btwnummer,iban,smtp_host,smtp_port,smtp_user,smtp_password,smtp_from,smtp_use_tls,app_base_url,mollie_key_test) VALUES (1,?,?,?,?,?,?,?,?,?,?,?,?)",
            (_ENV_NAAM, _ENV_KVK, _ENV_BTW, _ENV_IBAN, _ENV_SMTP_HOST, _ENV_SMTP_PORT,
             _ENV_SMTP_USER, _ENV_SMTP_PASS, _ENV_SMTP_FROM, 1 if _ENV_SMTP_TLS else 0,
             _ENV_BASE_URL, _ENV_MOLLIE),
        )
        db.commit()
    db.close()
    _maak_default_template()


# ─── Settings & context ───────────────────────────────────────────────────────

def get_settings():
    if "settings" not in g:
        row = get_db().execute("SELECT * FROM instellingen WHERE id=1").fetchone()
        g.settings = dict(row) if row else {}
    return g.settings


@app.context_processor
def inject_globals():
    try:
        s = get_settings()
    except Exception:
        s = {}
    try:
        db = get_db()
        fac_n = db.execute("SELECT COUNT(*) FROM facturen").fetchone()[0]
        kla_n = db.execute("SELECT COUNT(*) FROM klanten").fetchone()[0]
    except Exception:
        fac_n = kla_n = 0
    return {
        "settings": s,
        "light_mode": request.cookies.get("theme", "dark") == "light",
        "huidig_jaar": datetime.now().year,
        "fac_count": fac_n,
        "kla_count": kla_n,
        "vandaag": date.today().isoformat(),
    }


# ─── Calculation helpers ──────────────────────────────────────────────────────

def factuur_berekening(regels, korting=0.0, korting_type="pct"):
    subtotaal = sum(r["aantal"] * r["prijs_per_stuk"] for r in regels)
    korting = float(korting or 0)
    if korting_type == "pct":
        korting_bedrag = subtotaal * korting / 100
    else:
        korting_bedrag = min(korting, subtotaal)
    netto = subtotaal - korting_bedrag
    btw_per_tarief = {}
    for r in regels:
        bedrag = r["aantal"] * r["prijs_per_stuk"]
        ratio = bedrag / subtotaal if subtotaal > 0 else 0
        net_r = bedrag - korting_bedrag * ratio
        t = r["btw_percentage"]
        btw_per_tarief[t] = btw_per_tarief.get(t, 0) + net_r * t / 100
    btw_totaal = sum(btw_per_tarief.values())
    return {
        "subtotaal": round(subtotaal, 2),
        "korting_bedrag": round(korting_bedrag, 2),
        "netto": round(netto, 2),
        "btw_per_tarief": {k: round(v, 2) for k, v in btw_per_tarief.items()},
        "btw_totaal": round(btw_totaal, 2),
        "totaal": round(netto + btw_totaal, 2),
    }


def smtp_ok():
    s = get_settings()
    return bool(s.get("smtp_host") and s.get("smtp_user") and s.get("smtp_password"))


def verstuur_email(naar, onderwerp, html_body):
    s = get_settings()
    if not smtp_ok():
        raise RuntimeError("SMTP niet geconfigureerd. Stel in via Instellingen.")
    msg = MIMEMultipart("alternative")
    msg["Subject"] = onderwerp
    msg["From"] = s.get("smtp_from") or s["smtp_user"]
    msg["To"] = naar
    msg.attach(MIMEText(html_body, "html"))
    with smtplib.SMTP(s["smtp_host"], int(s.get("smtp_port", 587)), timeout=20) as srv:
        if s.get("smtp_use_tls", 1):
            srv.starttls()
        srv.login(s["smtp_user"], s["smtp_password"])
        srv.sendmail(msg["From"], [naar], msg.as_string())


def _mollie_key():
    s = get_settings()
    mode = s.get("mollie_mode", "test")
    key = s.get("mollie_key_live") if mode == "live" else s.get("mollie_key_test")
    return key or _ENV_MOLLIE


def maak_mollie_payment(factuur_id, totaal, factuurnummer):
    key = _mollie_key()
    if not key:
        raise RuntimeError("Geen Mollie API-sleutel ingesteld.")
    s = get_settings()
    base = (s.get("app_base_url") or _ENV_BASE_URL).rstrip("/")
    if not base:
        raise RuntimeError("Geen APP_BASE_URL ingesteld. Mollie heeft een publieke URL nodig.")
    db = get_db()
    f_row = db.execute(
        "SELECT k.naam FROM facturen f JOIN klanten k ON k.id=f.klant_id WHERE f.id=?",
        (factuur_id,)
    ).fetchone()
    klant_naam = f_row["naam"] if f_row else ""

    # Redirect URL: bevestigingspagina op it-bosch.nl of fallback naar app
    bevestiging_url = s.get("mollie_bevestiging_url", "").strip()
    if bevestiging_url:
        import urllib.parse as _up
        params = _up.urlencode({"nr": factuurnummer, "klant": klant_naam, "bedrag": f"{totaal:.2f}"})
        redirect_url = f"{bevestiging_url.rstrip('/')}/?{params}"
    else:
        redirect_url = f"{base}{url_for('factuur_bekijken', fid=factuur_id)}"

    # Webhook URL: relay op it-bosch.nl of fallback naar interne route
    relay_url = s.get("mollie_relay_url", "").strip()
    webhook_url = relay_url if relay_url else f"{base}{url_for('mollie_webhook')}"

    from mollie.api.client import Client
    mc = Client()
    mc.set_api_key(key)
    p = mc.payments.create({
        "amount": {"currency": "EUR", "value": f"{totaal:.2f}"},
        "description": f"Factuur {factuurnummer}",
        "redirectUrl": redirect_url,
        "webhookUrl": webhook_url,
        "metadata": {"factuur_id": factuur_id},
    })
    db = get_db()
    db.execute("UPDATE facturen SET betaallink=?, mollie_payment_id=? WHERE id=?",
               (p.checkout_url, p.id, factuur_id))
    db.commit()
    return p.checkout_url


def volgend_factuurnummer(db):
    row = db.execute("SELECT factuurprefix, factuurvolgend FROM instellingen WHERE id=1").fetchone()
    prefix = (row["factuurprefix"] or "FAC") if row else "FAC"
    volgnr = (row["factuurvolgend"] or 1) if row else 1
    return f"{prefix}-{volgnr:04d}"


def bump_volgnummer(db):
    db.execute("UPDATE instellingen SET factuurvolgend = factuurvolgend + 1 WHERE id=1")
    db.commit()


def volgend_periodiek_datum(d: date, interval: str) -> date:
    if interval == "wekelijks":
        return d + timedelta(weeks=1)
    if interval == "maandelijks":
        m = d.month % 12 + 1
        y = d.year + (1 if d.month == 12 else 0)
        return date(y, m, min(d.day, monthrange(y, m)[1]))
    if interval == "kwartaal":
        m = d.month + 3
        y = d.year + (m - 1) // 12
        m = (m - 1) % 12 + 1
        return date(y, m, min(d.day, monthrange(y, m)[1]))
    if interval == "jaarlijks":
        y = d.year + 1
        return date(y, d.month, min(d.day, monthrange(y, d.month)[1]))
    return d


def _reg_val(f, key):
    return f[key] if key in f.keys() else None


# ─── Word export ──────────────────────────────────────────────────────────────

def genereer_word_factuur(factuur, regels, bek, s):
    doc = DocxDoc()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.5)

    def p(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
        para = doc.add_paragraph()
        para.alignment = align
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return para

    p(s.get("bedrijfsnaam", ""), bold=True, size=16)
    adres_parts = [s.get("adres", ""), f"{s.get('postcode','')} {s.get('stad','')}".strip()]
    if s.get("kvk"):
        adres_parts.append(f"KvK: {s['kvk']}")
    if s.get("btwnummer"):
        adres_parts.append(f"BTW: {s['btwnummer']}")
    p("\n".join(filter(None, adres_parts)), size=10)

    doc.add_paragraph()

    h = p("FACTUUR", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.RIGHT, color=(0xF0, 0xA5, 0x00))
    details = [
        f"Nr: {factuur['factuurnummer']}",
        f"Datum: {factuur['factuurdatum']}",
        f"Vervaldatum: {factuur['vervaldatum']}",
    ]
    if _reg_val(factuur, "projectnummer"):
        details.append(f"Project: {factuur['projectnummer']}")
    p("\n".join(details), size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph("─" * 80)

    p("Factureren aan:", bold=True)
    klant_lines = [factuur["naam"]]
    if _reg_val(factuur, "adres"):
        klant_lines.append(factuur["adres"])
    postcity = f"{_reg_val(factuur,'postcode') or ''} {_reg_val(factuur,'plaats') or ''}".strip()
    if postcity:
        klant_lines.append(postcity)
    p("\n".join(klant_lines), size=11)
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    hrow = tbl.rows[0]
    for i, h in enumerate(["Omschrijving", "Aantal", "Eenheid", "Prijs excl.", "BTW%", "Totaal"]):
        c = hrow.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)

    for r in regels:
        row = tbl.add_row()
        vals = [
            r["omschrijving"],
            str(r["aantal"]).rstrip("0").rstrip(".") if "." in str(r["aantal"]) else str(int(r["aantal"])),
            r.get("eenheid", "st"),
            f"€ {r['prijs_per_stuk']:.2f}",
            f"{int(r['btw_percentage'])}%",
            f"€ {r['aantal'] * r['prijs_per_stuk']:.2f}",
        ]
        for i, v in enumerate(vals):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    def totaalrij(label, bedrag, bold=False, prefix=""):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"{label}: {prefix}€ {abs(bedrag):.2f}")
        run.bold = bold
        run.font.size = Pt(11 if bold else 10)

    totaalrij("Subtotaal excl. BTW", bek["subtotaal"])
    if bek["korting_bedrag"] > 0:
        totaalrij("Korting", bek["korting_bedrag"], prefix="-")
    for tarief, btw_bedrag in sorted(bek["btw_per_tarief"].items()):
        totaalrij(f"BTW {int(tarief)}%", btw_bedrag)
    totaalrij("TOTAAL INCL. BTW", bek["totaal"], bold=True)

    if _reg_val(factuur, "notities"):
        doc.add_paragraph()
        p("Notities:", bold=True)
        doc.add_paragraph(factuur["notities"])

    doc.add_paragraph()
    footer = " | ".join(filter(None, [
        f"IBAN: {s['iban']}" if s.get("iban") else None,
        f"BTW-nr: {s['btwnummer']}" if s.get("btwnummer") else None,
        f"KvK: {s['kvk']}" if s.get("kvk") else None,
    ]))
    if footer:
        para = doc.add_paragraph(footer)
        para.runs[0].font.size = Pt(9)

    voetnoot = s.get("juridisch_voetnoot") or "Bij overschrijding van de betalingstermijn bent u van rechtswege in verzuim en is de wettelijke handelsrente (art. 6:119a BW) verschuldigd."
    vn = doc.add_paragraph(voetnoot)
    vn.runs[0].font.size = Pt(8)
    vn.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


# ─── Word via docxtpl ────────────────────────────────────────────────────────

def _maak_default_template():
    """Create factuur_template.docx with {{placeholders}} if it doesn't exist yet."""
    if not HAS_DOCX or os.path.exists(TEMPLATE_PAD):
        return
    doc = DocxDoc()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(2.5)

    def para(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    # Bedrijfsblok
    para("{{bedrijfsnaam}}", bold=True, size=16, space_after=2)
    para("{{bedrijf_adres}}", size=10, space_after=0)
    para("{{bedrijf_postcode_stad}}", size=10, space_after=2)
    para("KvK: {{kvk}}  |  BTW: {{btwnummer}}", size=9, space_after=4)

    # Factuur-header rechts
    para("FACTUUR", bold=True, size=24, align=WD_ALIGN_PARAGRAPH.RIGHT, color=(0xF0, 0xA5, 0x00), space_after=2)
    para("Nr: {{factuurnummer}}", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=1)
    para("Datum: {{factuurdatum}}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=1)
    para("Vervaldatum: {{vervaldatum}}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=1)
    para("{% if projectnummer %}Project: {{projectnummer}}{% endif %}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)

    # Scheidingslijn
    para("─" * 90, size=8, space_after=4)

    # Klantblok
    para("Factureren aan:", bold=True, size=10, space_after=2)
    para("{{klant_naam}}", bold=True, size=12, space_after=1)
    para("{{klant_adres}}", size=10, space_after=1)
    para("{{klant_postcode_stad}}", size=10, space_after=1)
    para("{% if klant_btw_nummer %}BTW: {{klant_btw_nummer}}{% endif %}", size=9, space_after=8)

    # Regelstabel
    tbl = doc.add_table(rows=3, cols=6)
    tbl.style = "Table Grid"

    # Header row
    hdrs = ["Omschrijving", "Aantal", "Eenheid", "Prijs excl.", "BTW%", "Totaal"]
    for i, h in enumerate(hdrs):
        c = tbl.rows[0].cells[i]
        c.text = h
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = c._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '111111')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tc_pr.append(shd)

    # Template row (docxtpl row loop — {%tr for r in regels %})
    loop_row = tbl.rows[1]
    vals = [
        "{%tr for r in regels %}{{r.omschrijving}}",
        "{{r.aantal}}",
        "{{r.eenheid}}",
        "{{r.prijs}}",
        "{{r.btw_pct}}%",
        "{{r.totaal}}",
    ]
    for i, v in enumerate(vals):
        loop_row.cells[i].text = v
        loop_row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # End-loop row
    end_row = tbl.rows[2]
    end_row.cells[0].text = "{%tr endfor %}"
    for i in range(1, 6):
        end_row.cells[i].text = ""
    for cell in end_row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(1)

    doc.add_paragraph()

    # Totaalblok
    tot_tbl = doc.add_table(rows=4, cols=2)
    tot_tbl.style = "Table Grid"

    def tot_rij(row, label, waarde, bold=False):
        for i, txt in enumerate([label, waarde]):
            c = tot_tbl.rows[row].cells[i]
            c.text = txt
            run = c.paragraphs[0].runs[0]
            run.bold = bold
            run.font.size = Pt(10 if not bold else 12)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if i == 1 else WD_ALIGN_PARAGRAPH.LEFT

    tot_rij(0, "Subtotaal excl. BTW", "€ {{subtotaal}}")
    tot_rij(1, "{%tr for b in btw_regels %}BTW {{b.tarief}}%", "€ {{b.bedrag}}")
    tot_rij(2, "{%tr endfor %}", "")
    tot_rij(3, "TOTAAL INCL. BTW", "€ {{totaal}}", bold=True)

    doc.add_paragraph()

    # Notities (conditioneel)
    para("{% if has_notities %}Notities: {{notities}}{% endif %}", size=10, space_after=8)

    # Betaalgegevens
    para("─" * 90, size=8, space_after=4)
    para("{{betaalgegevens}}", size=9, space_after=2)
    para("{{voetnoot}}", size=8, color=(0x88, 0x88, 0x88), space_after=0)

    doc.save(TEMPLATE_PAD)


def genereer_word_via_template(factuur, regels, bek, s):
    """Fill factuur_template.docx with invoice data via docxtpl."""
    tpl = DocxTemplate(TEMPLATE_PAD)
    regels_ctx = []
    for r in regels:
        bedrag = r["aantal"] * r["prijs_per_stuk"]
        aantal_str = str(r["aantal"])
        pv = _reg_val(r, "periode_van") or ""
        pt = _reg_val(r, "periode_tot") or ""
        periode = f"{pv} t/m {pt}" if pv and pt else (pv or pt or "")
        regels_ctx.append({
            "omschrijving": r["omschrijving"],
            "aantal": aantal_str.rstrip("0").rstrip(".") if "." in aantal_str else str(int(r["aantal"])),
            "eenheid": (_reg_val(r, "eenheid") or "st"),
            "prijs": f"{r['prijs_per_stuk']:.2f}",
            "btw_pct": int(r["btw_percentage"]),
            "totaal": f"{bedrag:.2f}",
            "periode": periode,
            "heeft_periode": bool(periode),
        })
    btw_regels = [
        {"tarief": int(t), "bedrag": f"{b:.2f}"}
        for t, b in sorted(bek["btw_per_tarief"].items())
        if b > 0.001
    ]
    betaalinfo = " | ".join(filter(None, [
        f"IBAN: {s['iban']}" if s.get("iban") else None,
        f"BTW-nr: {s['btwnummer']}" if s.get("btwnummer") else None,
        f"KvK: {s['kvk']}" if s.get("kvk") else None,
    ]))
    ctx = {
        "bedrijfsnaam": s.get("bedrijfsnaam", ""),
        "bedrijf_adres": s.get("adres", ""),
        "bedrijf_postcode_stad": f"{s.get('postcode','')} {s.get('stad','')}".strip(),
        "kvk": s.get("kvk", ""),
        "btwnummer": s.get("btwnummer", ""),
        "iban": s.get("iban", ""),
        "factuurnummer": factuur["factuurnummer"],
        "factuurdatum": factuur["factuurdatum"],
        "vervaldatum": factuur["vervaldatum"],
        "projectnummer": _reg_val(factuur, "projectnummer") or "",
        "referentie": _reg_val(factuur, "referentie") or "",
        "klant_naam": factuur["naam"],
        "klant_adres": _reg_val(factuur, "adres") or "",
        "klant_postcode_stad": f"{_reg_val(factuur,'postcode') or ''} {_reg_val(factuur,'plaats') or ''}".strip(),
        "klant_btw_nummer": _reg_val(factuur, "btw_nummer") or "",
        "regels": regels_ctx,
        "subtotaal": f"{bek['subtotaal']:.2f}",
        "korting_bedrag": f"{bek['korting_bedrag']:.2f}",
        "has_korting": bek["korting_bedrag"] > 0,
        "btw_regels": btw_regels,
        "btw_totaal": f"{bek['btw_totaal']:.2f}",
        "totaal": f"{bek['totaal']:.2f}",
        "notities": _reg_val(factuur, "notities") or "",
        "has_notities": bool(_reg_val(factuur, "notities")),
        "betaalgegevens": betaalinfo,
        "voetnoot": s.get("juridisch_voetnoot") or "Bij overschrijding van de betalingstermijn is de wettelijke handelsrente (art. 6:119a BW) verschuldigd.",
    }
    tpl.render(ctx)
    return tpl


# ─── Thema ────────────────────────────────────────────────────────────────────

@app.route("/thema", methods=["POST"])
def thema_toggle():
    current = request.cookies.get("theme", "dark")
    resp = make_response(redirect(request.referrer or url_for("dashboard")))
    resp.set_cookie("theme", "light" if current == "dark" else "dark", max_age=31536000)
    return resp


# ─── Dashboard ────────────────────────────────────────────────────────────────

@app.route("/")
def dashboard():
    db = get_db()
    jaar = datetime.now().year
    open_f = db.execute(
        "SELECT COUNT(*) AS n, COALESCE(SUM((SELECT SUM(fr.aantal*fr.prijs_per_stuk*(1+fr.btw_percentage/100.0)) FROM factuurregels fr WHERE fr.factuur_id=facturen.id)),0) AS bedrag FROM facturen WHERE status IN ('open','verzonden')"
    ).fetchone()
    vervallen = db.execute(
        "SELECT COUNT(*) AS n FROM facturen WHERE status IN ('open','verzonden') AND vervaldatum < ?",
        (date.today().isoformat(),),
    ).fetchone()
    omzet = db.execute(
        "SELECT COALESCE(SUM(fr.aantal*fr.prijs_per_stuk),0) AS t FROM factuurregels fr JOIN facturen f ON f.id=fr.factuur_id WHERE f.factuurdatum LIKE ? AND f.status != 'concept'",
        (f"{jaar}-%",),
    ).fetchone()["t"]
    kosten = db.execute(
        "SELECT COALESCE(SUM(bedrag_excl_btw*(zakelijk_percentage/100.0)),0) AS t FROM uitgaven WHERE datum LIKE ?",
        (f"{jaar}-%",),
    ).fetchone()["t"]
    recente = db.execute(
        "SELECT f.*, k.naam AS klant_naam FROM facturen f JOIN klanten k ON k.id=f.klant_id ORDER BY f.id DESC LIMIT 8"
    ).fetchall()
    totalen = {}
    for f in recente:
        regels = db.execute("SELECT * FROM factuurregels WHERE factuur_id=?", (f["id"],)).fetchall()
        totalen[f["id"]] = factuur_berekening(regels, _reg_val(f, "korting") or 0, _reg_val(f, "korting_type") or "pct")["totaal"]
    return render_template("dashboard.html", open_f=open_f, vervallen=vervallen,
                           omzet=round(omzet, 2), kosten=round(kosten, 2),
                           winst=round(omzet - kosten, 2), recente=recente, totalen=totalen, jaar=jaar)


# ─── Klanten ──────────────────────────────────────────────────────────────────

@app.route("/klanten")
def klanten():
    db = get_db()
    alle = db.execute("SELECT * FROM klanten ORDER BY naam").fetchall()
    counts = {k["id"]: db.execute("SELECT COUNT(*) FROM facturen WHERE klant_id=?", (k["id"],)).fetchone()[0] for k in alle}
    return render_template("klanten.html", klanten=alle, counts=counts)


@app.route("/klanten/nieuw", methods=["GET", "POST"])
def klant_nieuw():
    if request.method == "POST":
        db = get_db()
        bt = request.form.get("betalingstermijn")
        db.execute(
            "INSERT INTO klanten (naam,contactpersoon,email,telefoon,adres,postcode,plaats,land,kvk,btw_nummer,betalingstermijn) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (request.form["naam"], request.form.get("contactpersoon"), request.form.get("email"),
             request.form.get("telefoon"), request.form.get("adres"), request.form.get("postcode"),
             request.form.get("plaats"), request.form.get("land", "Nederland"),
             request.form.get("kvk"), request.form.get("btw_nummer"),
             int(bt) if bt and bt.strip() else None),
        )
        db.commit()
        flash("Klant toegevoegd.", "success")
        return redirect(url_for("klanten"))
    return render_template("klant_form.html", klant=None)


@app.route("/klanten/<int:kid>/bewerken", methods=["GET", "POST"])
def klant_bewerken(kid):
    db = get_db()
    if request.method == "POST":
        bt = request.form.get("betalingstermijn")
        db.execute(
            "UPDATE klanten SET naam=?,contactpersoon=?,email=?,telefoon=?,adres=?,postcode=?,plaats=?,land=?,kvk=?,btw_nummer=?,betalingstermijn=? WHERE id=?",
            (request.form["naam"], request.form.get("contactpersoon"), request.form.get("email"),
             request.form.get("telefoon"), request.form.get("adres"), request.form.get("postcode"),
             request.form.get("plaats"), request.form.get("land", "Nederland"),
             request.form.get("kvk"), request.form.get("btw_nummer"),
             int(bt) if bt and bt.strip() else None, kid),
        )
        db.commit()
        flash("Klant bijgewerkt.", "success")
        return redirect(url_for("klanten"))
    return render_template("klant_form.html", klant=db.execute("SELECT * FROM klanten WHERE id=?", (kid,)).fetchone())


@app.route("/klanten/<int:kid>/verwijderen", methods=["POST"])
def klant_verwijderen(kid):
    db = get_db()
    db.execute("DELETE FROM klanten WHERE id=?", (kid,))
    db.commit()
    flash("Klant verwijderd.", "success")
    return redirect(url_for("klanten"))


# ─── Facturen ─────────────────────────────────────────────────────────────────

@app.route("/facturen")
def facturen():
    db = get_db()
    alle = db.execute("SELECT f.*, k.naam AS klant_naam FROM facturen f JOIN klanten k ON k.id=f.klant_id ORDER BY f.id DESC").fetchall()
    resultaten = []
    for f in alle:
        regels = db.execute("SELECT * FROM factuurregels WHERE factuur_id=?", (f["id"],)).fetchall()
        bek = factuur_berekening(regels, _reg_val(f, "korting") or 0, _reg_val(f, "korting_type") or "pct")
        resultaten.append({**dict(f), "totaal": bek["totaal"]})
    return render_template("facturen.html", facturen=resultaten)


def _sla_factuur_op(db, factuur_id=None):
    velden = (
        request.form["klant_id"], request.form["factuurdatum"], request.form["vervaldatum"],
        request.form.get("status", "concept"), request.form.get("notities"),
        float(request.form.get("korting", 0)), request.form.get("korting_type", "pct"),
        request.form.get("projectnummer"), request.form.get("referentie"),
        int(request.form.get("betalingstermijn", 30)),
    )
    if factuur_id:
        db.execute(
            "UPDATE facturen SET klant_id=?,factuurdatum=?,vervaldatum=?,status=?,notities=?,korting=?,korting_type=?,projectnummer=?,referentie=?,betalingstermijn=? WHERE id=?",
            (*velden, factuur_id),
        )
        db.execute("DELETE FROM factuurregels WHERE factuur_id=?", (factuur_id,))
    else:
        factuurnummer = request.form.get("factuurnummer") or volgend_factuurnummer(db)
        cur = db.execute(
            "INSERT INTO facturen (factuurnummer,klant_id,factuurdatum,vervaldatum,status,notities,korting,korting_type,projectnummer,referentie,betalingstermijn) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (factuurnummer, *velden),
        )
        factuur_id = cur.lastrowid
        bump_volgnummer(db)

    omschr_list = request.form.getlist("omschrijving")
    periode_van_list = request.form.getlist("periode_van")
    periode_tot_list = request.form.getlist("periode_tot")
    for i, (omschr, aantal, prijs, btw, eenheid) in enumerate(zip(
        omschr_list, request.form.getlist("aantal"),
        request.form.getlist("prijs_per_stuk"), request.form.getlist("btw_percentage"),
        request.form.getlist("eenheid"),
    )):
        if omschr.strip():
            pv = periode_van_list[i] if i < len(periode_van_list) else ""
            pt = periode_tot_list[i] if i < len(periode_tot_list) else ""
            db.execute(
                "INSERT INTO factuurregels (factuur_id,omschrijving,aantal,eenheid,prijs_per_stuk,btw_percentage,periode_van,periode_tot) VALUES (?,?,?,?,?,?,?,?)",
                (factuur_id, omschr, float(aantal or 1), eenheid or "st", float(prijs or 0), float(btw or 21),
                 pv or None, pt or None),
            )
    db.commit()
    return factuur_id


@app.route("/facturen/nieuw", methods=["GET", "POST"])
def factuur_nieuw():
    db = get_db()
    if request.method == "POST":
        fid = _sla_factuur_op(db)
        flash("Factuur aangemaakt.", "success")
        return redirect(url_for("factuur_bekijken", fid=fid))
    s = get_settings()
    return render_template("factuur_form.html",
                           klanten=db.execute("SELECT id,naam,betalingstermijn,adres,postcode,plaats,btw_nummer FROM klanten ORDER BY naam").fetchall(),
                           producten=db.execute("SELECT * FROM producten ORDER BY naam").fetchall(),
                           factuur=None, regels=[], voorgesteld_nummer=volgend_factuurnummer(db),
                           vandaag=date.today().isoformat(),
                           betalingstermijn=s.get("betalingstermijn", 30),
                           standaard_termijn=s.get("betalingstermijn", 30))


@app.route("/facturen/<int:fid>/bewerken", methods=["GET", "POST"])
def factuur_bewerken(fid):
    db = get_db()
    if request.method == "POST":
        _sla_factuur_op(db, factuur_id=fid)
        flash("Factuur bijgewerkt.", "success")
        return redirect(url_for("factuur_bekijken", fid=fid))
    factuur = db.execute("SELECT * FROM facturen WHERE id=?", (fid,)).fetchone()
    regels = db.execute("SELECT * FROM factuurregels WHERE factuur_id=?", (fid,)).fetchall()
    s = get_settings()
    return render_template("factuur_form.html",
                           klanten=db.execute("SELECT id,naam,betalingstermijn,adres,postcode,plaats,btw_nummer FROM klanten ORDER BY naam").fetchall(),
                           producten=db.execute("SELECT * FROM producten ORDER BY naam").fetchall(),
                           factuur=factuur, regels=regels,
                           vandaag=date.today().isoformat(),
                           betalingstermijn=_reg_val(factuur, "betalingstermijn") or 30,
                           standaard_termijn=s.get("betalingstermijn", 30),
                           voorgesteld_nummer=factuur["factuurnummer"])


@app.route("/facturen/<int:fid>")
def factuur_bekijken(fid):
    db = get_db()
    factuur = db.execute(
        "SELECT f.*, k.naam, k.adres, k.postcode, k.plaats, k.btw_nummer, k.email AS klant_email, f.id AS factuur_id FROM facturen f JOIN klanten k ON k.id=f.klant_id WHERE f.id=?",
        (fid,),
    ).fetchone()
    regels = db.execute("SELECT * FROM factuurregels WHERE factuur_id=?", (fid,)).fetchall()
    bek = factuur_berekening(regels, _reg_val(factuur, "korting") or 0, _reg_val(factuur, "korting_type") or "pct")
    return render_template("factuur_view.html", factuur=factuur, regels=regels, bek=bek,
                           mollie_actief=bool(_mollie_key()),
                           smtp_actief=smtp_ok(), has_docx=HAS_DOCX)


@app.route("/facturen/<int:fid>/status", methods=["POST"])
def factuur_status(fid):
    db = get_db()
    db.execute("UPDATE facturen SET status=? WHERE id=?", (request.form["status"], fid))
    db.commit()
    flash("Status bijgewerkt.", "success")
    return redirect(url_for("factuur_bekijken", fid=fid))


@app.route("/facturen/<int:fid>/dupliceren", methods=["POST"])
def factuur_dupliceren(fid):
    db = get_db()
    orig = db.execute("SELECT * FROM facturen WHERE id=?", (fid,)).fetchone()
    regels = db.execute("SELECT * FROM factuurregels WHERE factuur_id=?", (fid,)).fetchall()
    bt = _reg_val(orig, "betalingstermijn") or 30
    nr = volgend_factuurnummer(db)
    cur = db.execute(
        "INSERT INTO facturen (factuurnummer,klant_id,factuurdatum,vervaldatum,status,notities,korting,korting_type,projectnummer,referentie,betalingstermijn) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
        (nr, orig["klant_id"], date.today().isoformat(),
         (date.today() + timedelta(days=bt)).isoformat(), "concept",
         orig["notities"], _reg_val(orig, "korting") or 0,
         _reg_val(orig, "korting_type") or "pct",
         _reg_val(orig, "projectnummer"), _reg_val(orig, "referentie"), bt),
    )
    nid = cur.lastrowid
    for r in regels:
        db.execute(
            "INSERT INTO factuurregels (factuur_id,omschrijving,aantal,eenheid,prijs_per_stuk,btw_percentage) VALUES (?,?,?,?,?,?)",
            (nid, r["omschrijving"], r["aantal"], r.get("eenheid", "st"), r["prijs_per_stuk"], r["btw_percentage"]),
        )
    bump_volgnummer(db)
    db.commit()
    flash(f"Gedupliceerd als {nr}.", "success")
    return redirect(url_for("factuur_bewerken", fid=nid))


@app.route("/facturen/<int:fid>/word")
def factuur_word(fid):
    if not HAS_DOCX and not HAS_DOCXTPL:
        flash("python-docx niet geïnstalleerd.", "danger")
        return redirect(url_for("factuur_bekijken", fid=fid))
    db = get_db()
    factuur = db.execute(
        "SELECT f.*, k.naam, k.adres, k.postcode, k.plaats, k.btw_nummer FROM facturen f JOIN klanten k ON k.id=f.klant_id WHERE f.id=?",
        (fid,),
    ).fetchone()
    regels = db.execute("SELECT * FROM factuurregels WHERE factuur_id=?", (fid,)).fetchall()
    bek = factuur_berekening(regels, _reg_val(factuur, "korting") or 0, _reg_val(factuur, "korting_type") or "pct")
    buf = io.BytesIO()
    if HAS_DOCXTPL and os.path.exists(TEMPLATE_PAD):
        tpl = genereer_word_via_template(factuur, regels, bek, get_settings())
        tpl.save(buf)
    else:
        doc = genereer_word_factuur(factuur, regels, bek, get_settings())
        doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name=f"factuur_{factuur['factuurnummer']}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/facturen/<int:fid>/mollie-betaallink", methods=["POST"])
def factuur_mollie_betaallink(fid):
    db = get_db()
    regels = db.execute("SELECT * FROM factuurregels WHERE factuur_id=?", (fid,)).fetchall()
    factuur = db.execute("SELECT * FROM facturen WHERE id=?", (fid,)).fetchone()
    bek = factuur_berekening(regels, _reg_val(factuur, "korting") or 0, _reg_val(factuur, "korting_type") or "pct")
    try:
        maak_mollie_payment(fid, bek["totaal"], factuur["factuurnummer"])
        flash("Mollie betaallink aangemaakt.", "success")
    except Exception as e:
        flash(f"Mollie-fout: {e}", "danger")
    return redirect(url_for("factuur_bekijken", fid=fid))


@app.route("/facturen/<int:fid>/versturen", methods=["POST"])
def factuur_versturen(fid):
    db = get_db()
    factuur = db.execute(
        "SELECT f.*, k.naam AS klant_naam, k.email AS klant_email FROM facturen f JOIN klanten k ON k.id=f.klant_id WHERE f.id=?",
        (fid,),
    ).fetchone()
    regels = db.execute("SELECT * FROM factuurregels WHERE factuur_id=?", (fid,)).fetchall()
    bek = factuur_berekening(regels, _reg_val(factuur, "korting") or 0, _reg_val(factuur, "korting_type") or "pct")
    s = get_settings()
    if not factuur["klant_email"]:
        flash("Klant heeft geen e-mailadres.", "danger")
        return redirect(url_for("factuur_bekijken", fid=fid))
    betaallink = factuur["betaallink"]
    if not betaallink:
        try:
            betaallink = maak_mollie_payment(fid, bek["totaal"], factuur["factuurnummer"])
        except Exception:
            betaallink = None
    regels_html = "".join(
        f"<tr><td>{r['omschrijving']}</td><td>{r['aantal']} {r.get('eenheid','st')}</td><td>&euro; {r['prijs_per_stuk']:.2f}</td><td>&euro; {r['aantal']*r['prijs_per_stuk']:.2f}</td></tr>"
        for r in regels
    )
    betaal_html = f"<p><a href='{betaallink}'>Klik hier om direct online te betalen</a></p>" if betaallink else ""
    bedrijf = s.get("bedrijfsnaam", "")
    html = f"""<p>Beste {factuur['klant_naam']},</p>
<p>Hierbij ontvangt u factuur <strong>{factuur['factuurnummer']}</strong> van {bedrijf}, vervaldatum {factuur['vervaldatum']}.</p>
<table border="1" cellpadding="6" cellspacing="0"><tr><th>Omschrijving</th><th>Aantal</th><th>Prijs</th><th>Bedrag</th></tr>{regels_html}</table>
<p>Subtotaal: &euro; {bek['subtotaal']:.2f}<br>BTW: &euro; {bek['btw_totaal']:.2f}<br><strong>Totaal: &euro; {bek['totaal']:.2f}</strong></p>
{betaal_html}<p>Met vriendelijke groet,<br>{bedrijf}</p>"""
    try:
        verstuur_email(factuur["klant_email"], f"Factuur {factuur['factuurnummer']}", html)
        db.execute("UPDATE facturen SET status='verzonden' WHERE id=? AND status='concept'", (fid,))
        db.commit()
        flash(f"Factuur verstuurd naar {factuur['klant_email']}.", "success")
    except Exception as e:
        flash(f"E-mail mislukt: {e}", "danger")
    return redirect(url_for("factuur_bekijken", fid=fid))


@app.route("/mollie/webhook", methods=["POST"])
def mollie_webhook():
    payment_id = request.form.get("id")
    key = _mollie_key()
    if not payment_id or not key:
        return "", 400
    from mollie.api.client import Client
    mc = Client()
    mc.set_api_key(key)
    try:
        p = mc.payments.get(payment_id)
    except Exception:
        return "", 404
    fid = (p.metadata or {}).get("factuur_id")
    if not fid:
        return "", 200
    db = get_db()
    if p.is_paid():
        db.execute("UPDATE facturen SET status='betaald' WHERE id=?", (fid,))
    elif p.is_expired() or p.is_canceled() or p.is_failed():
        db.execute("UPDATE facturen SET status='vervallen' WHERE id=?", (fid,))
    db.commit()
    return "", 200


@app.route("/facturen/<int:fid>/verwijderen", methods=["POST"])
def factuur_verwijderen(fid):
    db = get_db()
    db.execute("DELETE FROM facturen WHERE id=?", (fid,))
    db.commit()
    flash("Factuur verwijderd.", "success")
    return redirect(url_for("facturen"))


# ─── Producten ────────────────────────────────────────────────────────────────

@app.route("/producten")
def producten():
    return render_template("producten.html", producten=get_db().execute("SELECT * FROM producten ORDER BY naam").fetchall())


@app.route("/producten/json")
def producten_json():
    rows = get_db().execute("SELECT id,naam,beschrijving,prijs,eenheid FROM producten ORDER BY naam").fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/producten/nieuw", methods=["GET", "POST"])
def product_nieuw():
    if request.method == "POST":
        db = get_db()
        db.execute("INSERT INTO producten (naam,beschrijving,prijs,eenheid) VALUES (?,?,?,?)",
                   (request.form["naam"], request.form.get("beschrijving"),
                    float(request.form.get("prijs", 0)), request.form.get("eenheid", "st")))
        db.commit()
        flash("Product toegevoegd.", "success")
        return redirect(url_for("producten"))
    return render_template("product_form.html", product=None)


@app.route("/producten/<int:pid>/bewerken", methods=["GET", "POST"])
def product_bewerken(pid):
    db = get_db()
    if request.method == "POST":
        db.execute("UPDATE producten SET naam=?,beschrijving=?,prijs=?,eenheid=? WHERE id=?",
                   (request.form["naam"], request.form.get("beschrijving"),
                    float(request.form.get("prijs", 0)), request.form.get("eenheid", "st"), pid))
        db.commit()
        flash("Product bijgewerkt.", "success")
        return redirect(url_for("producten"))
    return render_template("product_form.html", product=db.execute("SELECT * FROM producten WHERE id=?", (pid,)).fetchone())


@app.route("/producten/<int:pid>/verwijderen", methods=["POST"])
def product_verwijderen(pid):
    db = get_db()
    db.execute("DELETE FROM producten WHERE id=?", (pid,))
    db.commit()
    flash("Product verwijderd.", "success")
    return redirect(url_for("producten"))


# ─── Uitgaven ─────────────────────────────────────────────────────────────────

@app.route("/uitgaven")
def uitgaven():
    return render_template("uitgaven.html", uitgaven=get_db().execute("SELECT * FROM uitgaven ORDER BY datum DESC").fetchall())


@app.route("/uitgaven/nieuw", methods=["GET", "POST"])
def uitgave_nieuw():
    if request.method == "POST":
        db = get_db()
        db.execute(
            "INSERT INTO uitgaven (datum,omschrijving,categorie,bedrag_excl_btw,btw_percentage,zakelijk_percentage,leverancier,notities) VALUES (?,?,?,?,?,?,?,?)",
            (request.form["datum"], request.form["omschrijving"], request.form.get("categorie"),
             float(request.form["bedrag_excl_btw"]), float(request.form.get("btw_percentage", 21)),
             float(request.form.get("zakelijk_percentage", 100)),
             request.form.get("leverancier"), request.form.get("notities")),
        )
        db.commit()
        flash("Uitgave toegevoegd.", "success")
        return redirect(url_for("uitgaven"))
    return render_template("uitgave_form.html", vandaag=date.today().isoformat(), uitgave=None)


@app.route("/uitgaven/<int:uid>/bewerken", methods=["GET", "POST"])
def uitgave_bewerken(uid):
    db = get_db()
    if request.method == "POST":
        db.execute(
            "UPDATE uitgaven SET datum=?,omschrijving=?,categorie=?,bedrag_excl_btw=?,btw_percentage=?,zakelijk_percentage=?,leverancier=?,notities=? WHERE id=?",
            (request.form["datum"], request.form["omschrijving"], request.form.get("categorie"),
             float(request.form["bedrag_excl_btw"]), float(request.form.get("btw_percentage", 21)),
             float(request.form.get("zakelijk_percentage", 100)),
             request.form.get("leverancier"), request.form.get("notities"), uid),
        )
        db.commit()
        flash("Uitgave bijgewerkt.", "success")
        return redirect(url_for("uitgaven"))
    return render_template("uitgave_form.html", vandaag=date.today().isoformat(),
                           uitgave=db.execute("SELECT * FROM uitgaven WHERE id=?", (uid,)).fetchone())


@app.route("/uitgaven/<int:uid>/verwijderen", methods=["POST"])
def uitgave_verwijderen(uid):
    get_db().execute("DELETE FROM uitgaven WHERE id=?", (uid,))
    get_db().commit()
    flash("Uitgave verwijderd.", "success")
    return redirect(url_for("uitgaven"))


# ─── BTW ──────────────────────────────────────────────────────────────────────

@app.route("/btw")
def btw_overzicht():
    db = get_db()
    jaar = int(request.args.get("jaar", datetime.now().year))
    kwartalen = []
    for q in range(1, 5):
        sm = (q - 1) * 3 + 1
        start = f"{jaar}-{sm:02d}-01"
        eind = f"{jaar}-{sm+2:02d}-31"
        btw_v = db.execute("SELECT COALESCE(SUM(fr.aantal*fr.prijs_per_stuk*fr.btw_percentage/100.0),0) AS t FROM factuurregels fr JOIN facturen f ON f.id=fr.factuur_id WHERE f.factuurdatum BETWEEN ? AND ? AND f.status!='concept'", (start, eind)).fetchone()["t"]
        omzet = db.execute("SELECT COALESCE(SUM(fr.aantal*fr.prijs_per_stuk),0) AS t FROM factuurregels fr JOIN facturen f ON f.id=fr.factuur_id WHERE f.factuurdatum BETWEEN ? AND ? AND f.status!='concept'", (start, eind)).fetchone()["t"]
        btw_vb = db.execute("SELECT COALESCE(SUM(bedrag_excl_btw*btw_percentage/100.0*(zakelijk_percentage/100.0)),0) AS t FROM uitgaven WHERE datum BETWEEN ? AND ?", (start, eind)).fetchone()["t"]
        kosten = db.execute("SELECT COALESCE(SUM(bedrag_excl_btw*(zakelijk_percentage/100.0)),0) AS t FROM uitgaven WHERE datum BETWEEN ? AND ?", (start, eind)).fetchone()["t"]
        kwartalen.append({"kwartaal": q, "omzet_excl": round(omzet, 2), "kosten_excl": round(kosten, 2),
                          "btw_verschuldigd": round(btw_v, 2), "btw_voorbelasting": round(btw_vb, 2),
                          "btw_te_betalen": round(btw_v - btw_vb, 2)})
    return render_template("btw.html", kwartalen=kwartalen, jaar=jaar)


# ─── Jaaroverzicht ────────────────────────────────────────────────────────────

@app.route("/jaaroverzicht")
def jaaroverzicht():
    db = get_db()
    jaar = int(request.args.get("jaar", datetime.now().year))
    omzet = db.execute("SELECT COALESCE(SUM(fr.aantal*fr.prijs_per_stuk),0) AS t FROM factuurregels fr JOIN facturen f ON f.id=fr.factuur_id WHERE f.factuurdatum LIKE ? AND f.status!='concept'", (f"{jaar}-%",)).fetchone()["t"]
    kosten_cat = db.execute("SELECT COALESCE(categorie,'Overig') AS categorie, COALESCE(SUM(bedrag_excl_btw*(zakelijk_percentage/100.0)),0) AS totaal FROM uitgaven WHERE datum LIKE ? GROUP BY categorie ORDER BY totaal DESC", (f"{jaar}-%",)).fetchall()
    totale_kosten = sum(r["totaal"] for r in kosten_cat)
    return render_template("jaaroverzicht.html", jaar=jaar, omzet=round(omzet, 2),
                           kosten_per_cat=kosten_cat, totale_kosten=round(totale_kosten, 2),
                           winst=round(omzet - totale_kosten, 2))


# ─── Instellingen ─────────────────────────────────────────────────────────────

@app.route("/instellingen", methods=["GET", "POST"])
def instellingen():
    db = get_db()
    if request.method == "POST":
        f = request.form
        # Behoud bestaande waarden als het veld leeg wordt ingestuurd
        huidig = get_settings()
        def behoud(key, fallback=""):
            v = f.get(key, "").strip()
            return v if v else huidig.get(key, fallback)

        db.execute(
            """UPDATE instellingen SET bedrijfsnaam=?,adres=?,postcode=?,stad=?,email=?,telefoon=?,
               kvk=?,btwnummer=?,iban=?,betalingstermijn=?,factuurprefix=?,factuurvolgend=?,
               alg_voorwaarden=?,juridisch_voetnoot=?,mollie_key_test=?,mollie_key_live=?,mollie_mode=?,
               smtp_host=?,smtp_port=?,smtp_user=?,smtp_password=?,smtp_from=?,smtp_use_tls=?,app_base_url=?,
               mollie_poll_url=?,mollie_poll_token=?,mollie_relay_url=?,mollie_bevestiging_url=?
               WHERE id=1""",
            (f.get("bedrijfsnaam",""), f.get("adres",""), f.get("postcode",""), f.get("stad",""),
             f.get("email",""), f.get("telefoon",""), f.get("kvk",""), f.get("btwnummer",""),
             f.get("iban",""), int(f.get("betalingstermijn",30)),
             f.get("factuurprefix","FAC"), int(f.get("factuurvolgend",1)),
             f.get("alg_voorwaarden",""), f.get("juridisch_voetnoot",""),
             behoud("mollie_key_test"), behoud("mollie_key_live"), f.get("mollie_mode","test"),
             f.get("smtp_host",""), int(f.get("smtp_port",587)),
             f.get("smtp_user",""), behoud("smtp_password"), f.get("smtp_from",""),
             1 if f.get("smtp_use_tls") else 0, f.get("app_base_url",""),
             f.get("mollie_poll_url",""), behoud("mollie_poll_token"),
             f.get("mollie_relay_url",""), f.get("mollie_bevestiging_url",""))
        )
        db.commit()
        g.pop("settings", None)
        flash("Instellingen opgeslagen.", "success")
        return redirect(url_for("instellingen", tab=f.get("active_tab", "bedrijf")))
    s = get_settings()
    tab = request.args.get("tab", "bedrijf")
    compliance = {k: bool(s.get(k)) for k in ("bedrijfsnaam", "adres", "kvk", "btwnummer", "iban", "email")}
    return render_template("instellingen.html", s=s, tab=tab, compliance=compliance)


# ─── Word template beheer ─────────────────────────────────────────────────────

@app.route("/instellingen/word-template")
def word_template_download():
    if not os.path.exists(TEMPLATE_PAD):
        _maak_default_template()
    if not os.path.exists(TEMPLATE_PAD):
        flash("Geen Word-template beschikbaar (python-docx niet geïnstalleerd).", "danger")
        return redirect(url_for("instellingen"))
    return send_file(TEMPLATE_PAD, as_attachment=True, download_name="factuur_template.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/instellingen/word-template/upload", methods=["POST"])
def word_template_upload():
    f = request.files.get("template")
    if not f or not f.filename.endswith(".docx"):
        flash("Upload een geldig .docx bestand.", "danger")
        return redirect(url_for("instellingen"))
    f.save(TEMPLATE_PAD)
    flash("Word-template geüpload. Vanaf nu worden facturen via dit template gegenereerd.", "success")
    return redirect(url_for("instellingen"))


@app.route("/instellingen/word-template/reset", methods=["POST"])
def word_template_reset():
    if os.path.exists(TEMPLATE_PAD):
        os.remove(TEMPLATE_PAD)
    _maak_default_template()
    flash("Word-template teruggezet naar standaard.", "success")
    return redirect(url_for("instellingen"))


# ─── Periodieke facturen ──────────────────────────────────────────────────────

@app.route("/periodiek")
def periodieke_facturen():
    rows = get_db().execute("SELECT p.*, k.naam AS klant_naam FROM periodieke_facturen p JOIN klanten k ON k.id=p.klant_id ORDER BY p.volgende_datum").fetchall()
    return render_template("periodieke_facturen.html", schema_lijst=rows)


@app.route("/periodiek/nieuw", methods=["GET", "POST"])
def periodiek_nieuw():
    db = get_db()
    if request.method == "POST":
        omschr = request.form["omschrijving"]
        cur = db.execute(
            "INSERT INTO periodieke_facturen (naam,klant_id,interval,volgende_datum,actief,betalingstermijn) VALUES (?,?,?,?,1,?)",
            (omschr, request.form["klant_id"], request.form["interval"],
             request.form["volgende_datum"], int(request.form.get("betalingstermijn", 30))),
        )
        pid = cur.lastrowid
        # Main regel from bedrag_excl_btw
        bedrag = float(request.form.get("bedrag_excl_btw") or 0)
        btw = float(request.form.get("btw_percentage") or 21)
        pv = request.form.get("periode_van") or None
        pt = request.form.get("periode_tot") or None
        db.execute("INSERT INTO periodieke_regels (periodiek_id,omschrijving,aantal,eenheid,prijs_per_stuk,btw_percentage,periode_van,periode_tot) VALUES (?,?,?,?,?,?,?,?)",
                   (pid, omschr, 1, "st", bedrag, btw, pv, pt))
        # Extra regels
        for ro, ra, rp, rb, re, rpv, rpt in zip(
            request.form.getlist("regel_omschrijving"), request.form.getlist("regel_aantal"),
            request.form.getlist("regel_prijs"), request.form.getlist("regel_btw"),
            request.form.getlist("regel_eenheid"),
            request.form.getlist("regel_periode_van"),
            request.form.getlist("regel_periode_tot"),
        ):
            if ro.strip():
                db.execute("INSERT INTO periodieke_regels (periodiek_id,omschrijving,aantal,eenheid,prijs_per_stuk,btw_percentage,periode_van,periode_tot) VALUES (?,?,?,?,?,?,?,?)",
                           (pid, ro, float(ra or 1), re or "st", float(rp or 0), float(rb or 21), rpv or None, rpt or None))
        db.commit()
        flash("Periodieke factuur aangemaakt.", "success")
        return redirect(url_for("periodieke_facturen"))
    return render_template("periodiek_form.html",
                           klanten=db.execute("SELECT * FROM klanten ORDER BY naam").fetchall(),
                           schema=None, regels=[])


@app.route("/periodiek/<int:pid>/bewerken", methods=["GET", "POST"])
def periodiek_bewerken(pid):
    db = get_db()
    if request.method == "POST":
        omschr = request.form["omschrijving"]
        db.execute(
            "UPDATE periodieke_facturen SET naam=?,klant_id=?,interval=?,volgende_datum=?,betalingstermijn=? WHERE id=?",
            (omschr, request.form["klant_id"], request.form["interval"],
             request.form["volgende_datum"], int(request.form.get("betalingstermijn", 30)), pid),
        )
        db.execute("DELETE FROM periodieke_regels WHERE periodiek_id=?", (pid,))
        bedrag = float(request.form.get("bedrag_excl_btw") or 0)
        btw = float(request.form.get("btw_percentage") or 21)
        pv = request.form.get("periode_van") or None
        pt = request.form.get("periode_tot") or None
        db.execute("INSERT INTO periodieke_regels (periodiek_id,omschrijving,aantal,eenheid,prijs_per_stuk,btw_percentage,periode_van,periode_tot) VALUES (?,?,?,?,?,?,?,?)",
                   (pid, omschr, 1, "st", bedrag, btw, pv, pt))
        for ro, ra, rp, rb, re, rpv, rpt in zip(
            request.form.getlist("regel_omschrijving"), request.form.getlist("regel_aantal"),
            request.form.getlist("regel_prijs"), request.form.getlist("regel_btw"),
            request.form.getlist("regel_eenheid"),
            request.form.getlist("regel_periode_van"),
            request.form.getlist("regel_periode_tot"),
        ):
            if ro.strip():
                db.execute("INSERT INTO periodieke_regels (periodiek_id,omschrijving,aantal,eenheid,prijs_per_stuk,btw_percentage,periode_van,periode_tot) VALUES (?,?,?,?,?,?,?,?)",
                           (pid, ro, float(ra or 1), re or "st", float(rp or 0), float(rb or 21), rpv or None, rpt or None))
        db.commit()
        flash("Periodieke factuur bijgewerkt.", "success")
        return redirect(url_for("periodieke_facturen"))
    schema = db.execute("SELECT * FROM periodieke_facturen WHERE id=?", (pid,)).fetchone()
    regels = db.execute("SELECT * FROM periodieke_regels WHERE periodiek_id=?", (pid,)).fetchall()
    return render_template("periodiek_form.html",
                           klanten=db.execute("SELECT * FROM klanten ORDER BY naam").fetchall(),
                           schema=schema, regels=regels)


@app.route("/periodiek/<int:pid>/verwijderen", methods=["POST"])
def periodiek_verwijderen(pid):
    db = get_db()
    db.execute("DELETE FROM periodieke_facturen WHERE id=?", (pid,))
    db.commit()
    flash("Periodieke factuur verwijderd.", "success")
    return redirect(url_for("periodieke_facturen"))


@app.route("/periodiek/<int:pid>/toggle", methods=["POST"])
def periodiek_toggle(pid):
    db = get_db()
    db.execute("UPDATE periodieke_facturen SET actief = CASE WHEN actief=1 THEN 0 ELSE 1 END WHERE id=?", (pid,))
    db.commit()
    return redirect(url_for("periodieke_facturen"))


@app.route("/cron/periodiek", methods=["GET", "POST"])
def cron_periodiek():
    n = _verwerk_periodieke_facturen()
    flash(f"{n} periodieke factuur(en) verwerkt.", "success")
    return redirect(url_for("periodieke_facturen"))


@app.route("/mollie/check", methods=["POST"])
def mollie_check():
    fid = request.form.get("fid")
    s = get_settings()
    poll_url = s.get("mollie_poll_url", "").strip()
    poll_token = s.get("mollie_poll_token", "").strip()
    if not poll_url or not poll_token:
        flash("Poll URL of token niet ingesteld. Configureer dit via Instellingen → Mollie.", "danger")
    else:
        resultaat = _poll_mollie_status()
        categorie = "danger" if any(w in resultaat for w in ["fout", "Fout", "bereiken", "geüpload"]) else "success"
        flash(resultaat, categorie)
    if fid:
        return redirect(url_for("factuur_bekijken", fid=fid))
    return redirect(url_for("facturen"))


def _verwerk_periodieke_facturen():
    db = sqlite3.connect(DB_PATH)
    db.row_factory = sqlite3.Row
    try:
        vandaag = date.today().isoformat()
        due = db.execute("SELECT * FROM periodieke_facturen WHERE actief=1 AND volgende_datum<=?", (vandaag,)).fetchall()
        count = 0
        for p in due:
            regels = db.execute("SELECT * FROM periodieke_regels WHERE periodiek_id=?", (p["id"],)).fetchall()
            s_row = db.execute("SELECT factuurprefix,factuurvolgend FROM instellingen WHERE id=1").fetchone()
            prefix = (s_row["factuurprefix"] or "FAC") if s_row else "FAC"
            volgnr = (s_row["factuurvolgend"] or 1) if s_row else 1
            nr = f"{prefix}-{volgnr:04d}"
            bt = p["betalingstermijn"] or 30
            fd = date.today()
            vd = fd + timedelta(days=int(bt))
            cur = db.execute(
                "INSERT INTO facturen (factuurnummer,klant_id,factuurdatum,vervaldatum,status,notities,betalingstermijn) VALUES (?,?,?,?,?,?,?)",
                (nr, p["klant_id"], fd.isoformat(), vd.isoformat(), "open", p["notities"], bt),
            )
            fid = cur.lastrowid
            # Bereken periode op basis van interval en volgende_datum
        schema_datum = date.fromisoformat(p["volgende_datum"])
        interval = p["interval"]
        if interval == "maandelijks":
            pv = schema_datum.replace(day=1)
            pt = schema_datum.replace(day=monthrange(schema_datum.year, schema_datum.month)[1])
        elif interval == "kwartaal":
            q_start = ((schema_datum.month - 1) // 3) * 3 + 1
            pv = date(schema_datum.year, q_start, 1)
            q_end_m = q_start + 2
            pt = date(schema_datum.year, q_end_m, monthrange(schema_datum.year, q_end_m)[1])
        elif interval == "jaarlijks":
            pv = date(schema_datum.year, 1, 1)
            pt = date(schema_datum.year, 12, 31)
        else:
            pv = pt = None

        for r in regels:
            regel_pv = r["periode_van"] if r["periode_van"] else (pv.isoformat() if pv else None)
            regel_pt = r["periode_tot"] if r["periode_tot"] else (pt.isoformat() if pt else None)
            db.execute(
                "INSERT INTO factuurregels (factuur_id,omschrijving,aantal,eenheid,prijs_per_stuk,btw_percentage,periode_van,periode_tot) VALUES (?,?,?,?,?,?,?,?)",
                (fid, r["omschrijving"], r["aantal"], r["eenheid"] or "st", r["prijs_per_stuk"], r["btw_percentage"], regel_pv, regel_pt),
            )
            volgende = volgend_periodiek_datum(date.fromisoformat(p["volgende_datum"]), p["interval"])
            db.execute("UPDATE periodieke_facturen SET volgende_datum=? WHERE id=?", (volgende.isoformat(), p["id"]))
            db.execute("UPDATE instellingen SET factuurvolgend=factuurvolgend+1 WHERE id=1")
            count += 1
        db.commit()
        return count
    finally:
        db.close()


# ─── Mollie polling (geen publieke webhook nodig) ─────────────────────────────

def _stuur_betaalbevestiging(factuur_id: int):
    try:
        s = get_settings()
        if not smtp_ok():
            return
        db = get_db()
        row = db.execute(
            """SELECT f.factuurnummer, f.factuurdatum, k.naam, k.email,
                      k.adres, k.postcode, k.plaats
               FROM facturen f JOIN klanten k ON k.id=f.klant_id
               WHERE f.id=?""", (factuur_id,)
        ).fetchone()
        if not row or not row["email"]:
            return
        regels = db.execute(
            "SELECT omschrijving, aantal, eenheid, prijs_per_stuk, btw_percentage FROM factuurregels WHERE factuur_id=?",
            (factuur_id,)
        ).fetchall()
        totaal = sum(r["aantal"] * r["prijs_per_stuk"] * (1 + r["btw_percentage"] / 100) for r in regels)
        subtotaal = sum(r["aantal"] * r["prijs_per_stuk"] for r in regels)
        btw = totaal - subtotaal

        regels_html = "".join(
            f"""<tr>
                  <td style="padding:8px 12px;border-bottom:1px solid #f0f0f0">{r['omschrijving']}</td>
                  <td style="padding:8px 12px;border-bottom:1px solid #f0f0f0;text-align:right">{r['aantal']:g} {r['eenheid']}</td>
                  <td style="padding:8px 12px;border-bottom:1px solid #f0f0f0;text-align:right;font-family:monospace">€ {r['aantal']*r['prijs_per_stuk']:.2f}</td>
                </tr>"""
            for r in regels
        )

        html = f"""<!DOCTYPE html>
<html lang="nl">
<head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"></head>
<body style="margin:0;padding:0;background:#f5f5f5;font-family:'Segoe UI',Arial,sans-serif">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#f5f5f5;padding:40px 0">
<tr><td align="center">
<table width="580" cellpadding="0" cellspacing="0" style="background:#ffffff;border-radius:4px;overflow:hidden">

  <!-- Header -->
  <tr><td style="background:#111111;padding:28px 36px">
    <span style="font-size:20px;font-weight:700;color:#f0a500;font-family:monospace;letter-spacing:0.05em">
      {s.get('bedrijfsnaam','Mijn Bedrijf')}
    </span>
  </td></tr>

  <!-- Titel -->
  <tr><td style="padding:32px 36px 0">
    <div style="font-size:22px;font-weight:700;color:#111;margin-bottom:8px">Betaling ontvangen</div>
    <div style="color:#666;font-size:14px;line-height:1.6">
      Beste {row['naam']},<br><br>
      Wij hebben uw betaling in goede orde ontvangen. Hieronder vindt u een overzicht.
    </div>
  </td></tr>

  <!-- Factuurinfo -->
  <tr><td style="padding:24px 36px 0">
    <table width="100%" cellpadding="0" cellspacing="0" style="background:#f9f9f9;border-radius:4px">
      <tr>
        <td style="padding:12px 16px;font-size:12px;color:#888;letter-spacing:0.06em;font-family:monospace">FACTUURNUMMER</td>
        <td style="padding:12px 16px;font-size:13px;font-weight:600;font-family:monospace;color:#f0a500">{row['factuurnummer']}</td>
      </tr>
      <tr style="border-top:1px solid #eeeeee">
        <td style="padding:12px 16px;font-size:12px;color:#888;letter-spacing:0.06em;font-family:monospace">FACTUURDATUM</td>
        <td style="padding:12px 16px;font-size:13px;font-family:monospace">{row['factuurdatum']}</td>
      </tr>
      <tr style="border-top:1px solid #eeeeee">
        <td style="padding:12px 16px;font-size:12px;color:#888;letter-spacing:0.06em;font-family:monospace">STATUS</td>
        <td style="padding:12px 16px"><span style="background:#e8f5e9;color:#2e7d32;font-size:11px;font-family:monospace;padding:3px 8px;border-radius:2px">BETAALD</span></td>
      </tr>
    </table>
  </td></tr>

  <!-- Regeloverzicht -->
  <tr><td style="padding:24px 36px 0">
    <table width="100%" cellpadding="0" cellspacing="0">
      <thead>
        <tr style="background:#111">
          <th style="padding:10px 12px;text-align:left;font-size:11px;color:#f0a500;font-family:monospace;letter-spacing:0.06em;font-weight:600">OMSCHRIJVING</th>
          <th style="padding:10px 12px;text-align:right;font-size:11px;color:#f0a500;font-family:monospace;letter-spacing:0.06em;font-weight:600">AANTAL</th>
          <th style="padding:10px 12px;text-align:right;font-size:11px;color:#f0a500;font-family:monospace;letter-spacing:0.06em;font-weight:600">BEDRAG</th>
        </tr>
      </thead>
      <tbody>{regels_html}</tbody>
    </table>
  </td></tr>

  <!-- Totalen -->
  <tr><td style="padding:0 36px 0">
    <table width="100%" cellpadding="0" cellspacing="0" style="border-top:2px solid #f0a500;margin-top:0">
      <tr>
        <td style="padding:10px 12px;font-size:13px;color:#666">Subtotaal excl. BTW</td>
        <td style="padding:10px 12px;text-align:right;font-family:monospace;font-size:13px">€ {subtotaal:.2f}</td>
      </tr>
      <tr>
        <td style="padding:10px 12px;font-size:13px;color:#666">BTW</td>
        <td style="padding:10px 12px;text-align:right;font-family:monospace;font-size:13px">€ {btw:.2f}</td>
      </tr>
      <tr style="background:#f9f9f9">
        <td style="padding:14px 12px;font-size:16px;font-weight:700">Totaal betaald</td>
        <td style="padding:14px 12px;text-align:right;font-family:monospace;font-size:16px;font-weight:700;color:#f0a500">€ {totaal:.2f}</td>
      </tr>
    </table>
  </td></tr>

  <!-- Footer -->
  <tr><td style="padding:32px 36px;border-top:1px solid #eeeeee;margin-top:24px">
    <div style="font-size:12px;color:#aaa;line-height:1.8">
      {s.get('bedrijfsnaam','')} {'• KvK: ' + s.get('kvk','') if s.get('kvk') else ''}
      {'• BTW: ' + s.get('btwnummer','') if s.get('btwnummer') else ''}<br>
      {'IBAN: ' + s.get('iban','') if s.get('iban') else ''}
    </div>
  </td></tr>

</table>
</td></tr></table>
</body></html>"""

        msg = MIMEMultipart("alternative")
        msg["Subject"] = f"Betaalbevestiging {row['factuurnummer']} — {s.get('bedrijfsnaam','')}"
        msg["From"] = s.get("smtp_from") or s.get("smtp_user", "")
        msg["To"] = row["email"]
        msg.attach(MIMEText(html, "html", "utf-8"))

        with smtplib.SMTP(s["smtp_host"], int(s.get("smtp_port", 587)), timeout=20) as srv:
            if s.get("smtp_use_tls", 1):
                srv.starttls()
            srv.login(s["smtp_user"], s["smtp_password"])
            srv.sendmail(msg["From"], [row["email"]], msg.as_string())
    except Exception:
        pass


def _verifieer_mollie_betaling(payment_id: str) -> bool:
    """Verifieert rechtstreeks bij Mollie of een betaling echt 'paid' is."""
    try:
        key = _mollie_key()
        if not key:
            return False
        req = urllib.request.Request(
            f"https://api.mollie.com/v2/payments/{urllib.parse.quote(payment_id)}",
            headers={"Authorization": f"Bearer {key}", "User-Agent": "ZZP-Boekhouding/2"},
        )
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = _json.loads(resp.read().decode())
        return data.get("status") == "paid"
    except Exception:
        return False


def _poll_mollie_status() -> str:
    """Pollt it-bosch.nl voor nieuwe betaalstatussen. Geeft statusbericht terug."""
    with app.app_context():
        try:
            s = get_settings()
            poll_url = s.get("mollie_poll_url", "").strip()
            poll_token = s.get("mollie_poll_token", "").strip()
            if not poll_url or not poll_token:
                return "Poll URL of token niet ingesteld."
            url = f"{poll_url.rstrip('/')}?token={urllib.parse.quote(poll_token)}"
            req = urllib.request.Request(url, headers={
                "User-Agent": "Mozilla/5.0 (compatible; ZZP-Boekhouding/2; +https://it-bosch.nl)",
                "Accept": "application/json",
            })
            try:
                with urllib.request.urlopen(req, timeout=10) as resp:
                    updates = _json.loads(resp.read().decode())
            except urllib.error.HTTPError as e:
                return f"Verbindingsfout met it-bosch.nl: HTTP {e.code}. Zijn de PHP-bestanden geüpload?"
            except Exception as e:
                return f"Kan it-bosch.nl niet bereiken: {e}"
            if not updates:
                return "Geen nieuwe betalingen gevonden."
            db = get_db()
            bijgewerkt = 0
            for upd in updates:
                pid = upd.get("payment_id", "")
                if not pid:
                    continue
                if not _verifieer_mollie_betaling(pid):
                    continue
                row = db.execute("SELECT id FROM facturen WHERE mollie_payment_id=?", (pid,)).fetchone()
                if row:
                    db.execute("UPDATE facturen SET status='betaald' WHERE id=?", (row["id"],))
                    _stuur_betaalbevestiging(row["id"])
                    bijgewerkt += 1
            db.commit()
            return f"{bijgewerkt} factuur/facturen bijgewerkt naar betaald." if bijgewerkt else "Geen nieuwe betalingen gevonden."
        except Exception as e:
            return f"Onverwachte fout: {e}"


# ─── Scheduler ────────────────────────────────────────────────────────────────

if HAS_SCHEDULER:
    import urllib.parse
    _scheduler = BackgroundScheduler()
    _scheduler.add_job(_verwerk_periodieke_facturen, "cron", hour=8, minute=0, id="periodiek_daily")
    _scheduler.add_job(_poll_mollie_status, "interval", minutes=5, id="mollie_poll")
    _scheduler.start()


# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", debug=os.environ.get("FLASK_DEBUG", "false").lower() == "true", port=PORT)
